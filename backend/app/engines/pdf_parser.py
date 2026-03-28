import io
import re
from dataclasses import dataclass, field

from app.utils.exceptions import (
    InvalidFileError,
    EncryptedPDFError,
    FileTooLargeError,
    EmptyDocumentError,
)
from app.utils.text_processing import clean_text
from app.utils.logger import get_logger

logger = get_logger(__name__)

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
PDF_MAGIC = b"%PDF"
DOCX_MAGIC = b"PK"


@dataclass
class PDFParseResult:
    text: str = ""
    page_count: int = 0
    word_count: int = 0
    has_tables: bool = False
    has_images: bool = False
    has_columns: bool = False
    warnings: list[str] = field(default_factory=list)


class PDFParser:
    def extract(self, file_bytes: bytes, filename: str = "") -> PDFParseResult:
        if len(file_bytes) > MAX_FILE_SIZE:
            raise FileTooLargeError(f"File is {len(file_bytes) / 1024 / 1024:.1f}MB, max is 10MB")

        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

        if ext == "docx" or (not ext and file_bytes[:2] == DOCX_MAGIC):
            return self._extract_docx(file_bytes)

        if ext == "pdf" or file_bytes[:4] == PDF_MAGIC:
            return self._extract_pdf(file_bytes)

        # Check magic bytes
        if file_bytes[:4] == PDF_MAGIC:
            return self._extract_pdf(file_bytes)
        if file_bytes[:2] == DOCX_MAGIC:
            return self._extract_docx(file_bytes)

        raise InvalidFileError(f"Unsupported file type: '{ext or 'unknown'}'. Upload a PDF or DOCX file.")

    def _extract_pdf(self, file_bytes: bytes) -> PDFParseResult:
        try:
            import pdfplumber
        except ImportError:
            raise InvalidFileError("PDF processing library not available")

        result = PDFParseResult()

        try:
            pdf = pdfplumber.open(io.BytesIO(file_bytes))
        except Exception as e:
            if "password" in str(e).lower() or "encrypted" in str(e).lower():
                raise EncryptedPDFError()
            raise InvalidFileError(f"Could not open PDF: {e}")

        try:
            pages = pdf.pages
            result.page_count = len(pages)
            all_text_parts: list[str] = []

            for page in pages:
                try:
                    text = page.extract_text() or ""
                    all_text_parts.append(text)

                    # Check for tables
                    tables = page.extract_tables()
                    if tables:
                        result.has_tables = True
                        for table in tables:
                            for row in table:
                                if row:
                                    row_text = " | ".join(str(cell) for cell in row if cell)
                                    all_text_parts.append(row_text)

                    # Check for images
                    if page.images:
                        result.has_images = True

                except Exception as e:
                    result.warnings.append(f"Error extracting page {pages.index(page) + 1}: {e}")
                    continue

            pdf.close()
        except Exception as e:
            pdf.close()
            raise InvalidFileError(f"Error processing PDF: {e}")

        raw_text = "\n".join(all_text_parts)
        result.text = self._clean_extracted_text(raw_text)
        result.word_count = len(result.text.split())
        result.has_columns = self._detect_columns(raw_text)

        if not result.text.strip():
            raise EmptyDocumentError("PDF has no extractable text. It may be a scanned image.")

        return result

    def _extract_docx(self, file_bytes: bytes) -> PDFParseResult:
        try:
            from docx import Document
        except ImportError:
            raise InvalidFileError("DOCX processing library not available")

        result = PDFParseResult()

        try:
            doc = Document(io.BytesIO(file_bytes))
        except Exception as e:
            raise InvalidFileError(f"Could not open DOCX: {e}")

        paragraphs: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract from tables too
        for table in doc.tables:
            result.has_tables = True
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        raw_text = "\n".join(paragraphs)
        result.text = self._clean_extracted_text(raw_text)
        result.word_count = len(result.text.split())
        result.page_count = 1  # Approximate

        if not result.text.strip():
            raise EmptyDocumentError("DOCX has no extractable text.")

        return result

    def _clean_extracted_text(self, text: str) -> str:
        # Remove page numbers
        text = re.sub(r"\n\s*(?:Page\s+)?\d+\s*(?:of\s+\d+)?\s*\n", "\n", text, flags=re.IGNORECASE)

        # Remove repeated header/footer lines
        lines = text.split("\n")
        if len(lines) > 20:
            first_lines = [l.strip() for l in lines[:3] if l.strip()]
            cleaned_lines: list[str] = []
            for i, line in enumerate(lines):
                if i > 2 and line.strip() in first_lines:
                    continue  # Likely repeated header
                cleaned_lines.append(line)
            text = "\n".join(cleaned_lines)

        text = clean_text(text)
        return text

    def _detect_columns(self, text: str) -> bool:
        column_lines = 0
        for line in text.split("\n"):
            if re.search(r"\S\s{4,}\S", line):
                column_lines += 1
        return column_lines > 5
